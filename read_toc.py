# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document('HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG CƠ SỞ TẠI THÀNH PHỐ HỒ CHÍ MINH.docx')

current_page = 1
toc_items = []

print("=" * 70)
print("MỤC LỤC VỚI SỐ TRANG")
print("=" * 70)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        # Đếm ký tự để ước tính trang (khoảng 3000 ký tự/trang A4)
        # Hoặc đếm paragraph
        
        # In ra các heading/chapter
        if any(x in text.upper() for x in ['CHƯƠNG', 'I.', 'II.', 'III.', 'IV.', 'V.', 'VI.', 'VII.', 'VIII.', 'MỤC LỤC', 'KẾT LUẬN', 'TÀI LIỆU']):
            print(f"\n{text}")
        elif len(text) < 100 and len(text) > 5:
            print(f"  {text[:80]}...")

print("\n" + "=" * 70)
print("TỔNG SỐ PARAGRAPH:", len(doc.paragraphs))
print("TỔNG SỐ BẢNG:", len(doc.tables))
print("=" * 70)
